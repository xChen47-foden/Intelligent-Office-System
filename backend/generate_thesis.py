from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.enum.section import WD_HEADER_FOOTER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as qn2

# 创建文档
doc = Document()

# 设置页眉
for section in doc.sections:
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = "浙江工贸职业技术学校毕业论文"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

# 设置中文字体
def set_font(run, name='宋体', size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold

# 封面
doc.add_paragraph('\n\n\n\n\n\n\n\n', style=None)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('基于大模型的智能办公应用软件')
set_font(run, size=22, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n')
run = p.add_run('作者：钱星晨\n')
set_font(run, size=14)
run = p.add_run('指导老师：管璐欣\n')
set_font(run, size=14)
run = p.add_run('学号：2312072431\n')
set_font(run, size=14)
run = p.add_run('人工智能技术应用专业\n')
set_font(run, size=14)
run = p.add_run('2024年6月\n')
set_font(run, size=14)

doc.add_page_break()

# 摘要
doc.add_heading('摘  要', level=1)
abs_text = (
    "本论文介绍了一款基于大模型技术的智能办公应用软件的设计与实现。系统采用Vue3+TypeScript+Element-Plus为前端技术栈，"
    "后端采用FastAPI与Node.js，集成了大模型（如GPT）能力，实现了智能文档处理、日程管理、知识库等功能。系统界面美观，操作便捷，"
    "能够有效提升办公效率。论文详细阐述了系统的需求分析、架构设计、关键技术实现及测试过程，并对未来的优化方向进行了展望。\n"
    "关键词：大模型，智能办公，Vue3，FastAPI，文档处理，知识库"
)
doc.add_paragraph(abs_text)
doc.add_page_break()

# Abstract
doc.add_heading('Abstract', level=1)
abs_en = (
    "This thesis focuses on the design and implementation of an intelligent office application based on large language models. "
    "The system adopts a front-end and back-end separation architecture, with the front-end based on Vue3, TypeScript, and Element-Plus, "
    "and the back-end using FastAPI and Node.js. It integrates large model capabilities (such as GPT) to realize intelligent document processing, "
    "schedule management, and knowledge base functions. The system features a beautiful interface and convenient operation, effectively improving office efficiency. "
    "The thesis elaborates on the requirement analysis, architecture design, key technology implementation, and testing process, and looks forward to future optimization directions.\n"
    "Keywords: Large Language Model, Intelligent Office, Vue3, FastAPI, Document Processing, Knowledge Base"
)
doc.add_paragraph(abs_en)
doc.add_page_break()

def add_toc(paragraph):
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn2('w:instr'), 'TOC \\o \"1-3\" \\h \\z \\u')
    paragraph._p.append(fldSimple)

# 目录（自动）
doc.add_heading('目  录', level=1)
toc_paragraph = doc.add_paragraph()
add_toc(toc_paragraph)
doc.add_page_break()

# 正文各章节
def add_section(title, content):
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    doc.add_paragraph('')

add_section('第一章 绪论',
    """1.1 研究背景与意义
随着人工智能技术的快速发展，基于大模型的智能办公应用逐渐成为提升企业办公效率的重要工具。传统办公软件在处理复杂任务时存在一定局限，而大模型具备强大的自然语言理解与生成能力，能够为办公自动化、智能化提供有力支撑。

1.2 国内外研究现状
目前，国内外众多企业和科研机构均在探索大模型在办公场景的应用，如微软Copilot、Google Workspace AI助手等。国内也有多家企业推出了基于大模型的智能办公产品，但大多聚焦于单一场景，缺乏一体化、可扩展的解决方案。

1.3 主要研究内容
本论文以"智能办公助手"项目为例，系统性介绍了基于大模型的智能办公应用的设计与实现，包括需求分析、系统架构、关键技术、功能实现、测试与优化等内容。
"""
)

add_section('第二章 相关技术介绍',
    "2.1 大模型技术\n"
    "本项目集成了大模型（如GPT）能力，主要用于文档智能摘要、知识问答等场景。通过Prompt工程、RAG（检索增强生成）、Agent等技术，实现了自然语言理解与生成，极大提升了办公自动化与智能化水平。\n\n"
    "2.2 前端技术\n"
    "前端采用Vue3、TypeScript、Element-Plus等主流技术，具备响应式设计、主题切换、组件化开发等优势。\n\n"
    "2.3 后端技术\n"
    "后端采用FastAPI（Python）和Node.js，分别负责主业务聚合和子服务扩展，支持RESTful API设计，易于维护和扩展。\n\n"
    "2.4 数据库与持久化\n"
    "系统支持多种数据库（如SQLite、Redis），实现数据的高效存储与检索。"
)

add_section('第三章 系统需求分析',
    "3.1 功能需求\n"
    "- 支持本地部署或云端服务\n"
    "- 可视化操作界面，良好用户体验\n"
    "- 文本内容录入与文档上传\n"
    "- 集成大模型，实现智能助手\n"
    "- 支持知识库内容检索\n"
    "- 智能日程规划与管理\n"
    "- 创新功能模块扩展\n\n"
    "3.2 非功能需求\n"
    "- 界面美观，布局合理\n"
    "- 响应速度快，稳定性高\n"
    "- 提供详细操作手册"
)

add_section('第四章 系统设计',
    "4.1 总体架构设计\n"
    "系统采用B/S架构，前后端分离，支持多端访问。前端负责用户交互与展示，后端负责业务逻辑与数据处理。\n"
    "架构图示意：用户 → 前端（Vue3）→ 后端（FastAPI/Node.js）→ 数据库/大模型API\n\n"
    "4.2 前端设计\n"
    "采用组件化开发，支持主题切换、多标签页、全局搜索等功能。\n\n"
    "4.3 后端设计\n"
    "主后端采用FastAPI，负责聚合与调度，Node.js子服务支持扩展。\n\n"
    "4.4 数据库设计\n"
    "采用SQLite存储文档、用户、知识库等数据，Redis用于缓存与验证码等场景。"
)

add_section('第五章 系统实现',
    "5.1 关键模块与代码实现\n"
    "（1）智能文档处理模块（前端代码片段）\n"
    "```typescript\n"
    "async function searchDocs() {\n"
    "  loading.value = true\n"
    "  try {\n"
    "    await fetchDocList()\n"
    "  } finally {\n"
    "    loading.value = false\n"
    "  }\n"
    "}\n"
    "async function getSummary(docId: number) {\n"
    "  const response = await axios.post('/api/doc/summary', { id: docId })\n"
    "  summaryContent.value = response.data.summary\n"
    "}\n"
    "```\n"
    "（2）后端智能摘要服务（Node.js代码片段）\n"
    "```js\n"
    "router.post('/summary', async (req, res) => {\n"
    "  const { id } = req.body;\n"
    "  db.get(`SELECT content FROM documents WHERE id = ?`, [id], async (err, row) => {\n"
    "    if (err || !row) return res.status(404).json({ error: '文档不存在' });\n"
    "    try {\n"
    "      const summary = await getSummary(row.content.slice(0, 2000));\n"
    "      res.json({ summary });\n"
    "    } catch (e) {\n"
    "      res.status(500).json({ error: '大模型调用失败', detail: e.message });\n"
    "    }\n"
    "  });\n"
    "});\n"
    "```\n"
    "（3）全局搜索功能（Vue3代码片段）\n"
    "```typescript\n"
    "const search = (val: string) => {\n"
    "  if (val) {\n"
    "    searchResult.value = flattenAndFilterMenuItems(menuList.value, val)\n"
    "  } else {\n"
    "    searchResult.value = []\n"
    "  }\n"
    "}\n"
    "```\n"
    "（4）数据持久化存储（TypeScript代码片段）\n"
    "```typescript\n"
    "function saveStoreStorage<T>(newData: T) {\n"
    "  const version = import.meta.env.VITE_VERSION\n"
    "  initVersion(version)\n"
    "  const vs = localStorage.getItem('version') || version\n"
    "  const storedData = JSON.parse(localStorage.getItem(`sys-v${vs}`) || '{}')\n"
    "  const mergedData = { ...storedData, ...newData }\n"
    "  localStorage.setItem(`sys-v${vs}`, JSON.stringify(mergedData))\n"
    "}\n"
    "```\n"
    "5.2 主要功能界面展示\n"
    "（此处可插入系统界面截图，如智能文档、日程管理、知识库等页面）"
)

add_section('第六章 系统测试',
    "6.1 测试环境\n"
    "- 操作系统：Windows 10/Loongnix\n"
    "- 浏览器：Chrome、Firefox等\n"
    "- 服务器：本地/云端\n\n"
    "6.2 测试用例与结果\n"
    "- 文档上传、检索、编辑、下载功能测试\n"
    "- 智能摘要功能测试\n"
    "- 日程管理与提醒功能测试\n"
    "- 知识库检索与管理测试\n"
    "（可插入测试截图和结果表格）"
)

add_section('第七章 总结与展望',
    "本文实现了一款基于大模型的智能办公应用软件，具备智能文档处理、日程管理、知识库等功能。系统架构清晰，界面美观，用户体验良好。未来可进一步集成更多AI能力，扩展更多办公场景，提升系统智能化水平。"
)

# 参考文献
doc.add_heading('参考文献', level=1)
doc.add_paragraph(
    "[1] OpenAI. GPT-3技术文档.\n"
    "[2] Vue.js官方文档.\n"
    "[3] FastAPI官方文档.\n"
    "[4] Element-Plus官方文档.\n"
    "[5] 钱星晨. 智能办公助手项目文档."
)
doc.add_page_break()

# 附录
doc.add_heading('附录', level=1)
doc.add_paragraph(
    "一、安装与部署说明\n"
    "```bash\n"
    "pnpm install\n"
    "pnpm dev\n"
    "pip install -r requirements.txt\n"
    "uvicorn main:app --reload --port 8000\n"
    "```\n"
    "二、操作手册\n"
    "- 登录系统，进入主界面\n"
    "- 上传文档，支持txt、docx、xlsx等格式\n"
    "- 检索、编辑、下载文档\n"
    "- 使用智能摘要、知识库等功能"
)

# 保存文档
doc.save('基于大模型的智能办公应用软件_钱星晨.docx')
print("论文已生成：基于大模型的智能办公应用软件_钱星晨.docx")
